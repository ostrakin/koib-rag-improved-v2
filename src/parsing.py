# -*- coding: utf-8 -*-
"""
Document parsing for KOIB RAG.

v4.11 changes:
- supports OCR-recognized DOCX, CSV, TXT/MD in addition to PDF;
- keeps DOCX paragraphs/tables in document order instead of extracting all tables first;
- prevents old JSONL index artifacts from being indexed as plain text;
- fixes over-aggressive formula detection that treated ordinary hyphenated text as formulas.
"""
from __future__ import annotations

import csv
import gc
import io
import logging
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

import fitz
from docx import Document as DocxDocument
from PIL import Image

try:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
except Exception:  # pragma: no cover - python-docx import fallback
    CT_Tbl = CT_P = DocxTable = Paragraph = None

from .utils import (
    clean_text,
    deep_clean_ocr_text,
    detect_model_from_filename,
    detect_model_in_text,
    extract_headings,
    find_figure_caption,
    looks_like_jsonl_artifact,
    normalize_ocr_text,
    text_hash,
    StickyModelTagger,
)
from config import (
    FIGURES_DIR,
    MIN_IMAGE_HEIGHT,
    MIN_IMAGE_WIDTH,
    OCR_DPI,
    OCR_MIN_TEXT_CHARS,
    rel_figure_path,
)
from .text_processing import (
    LATEX_BLOCK_RE as _LATEX_BLOCK_RE,
    LATEX_INLINE_RE as _LATEX_INLINE_RE,
    MATH_STRICT_RE as _MATH_STRICT_RE,
    is_supported_artifact as _is_supported_artifact,
    is_true_formula_text as _is_true_formula_text,
)

logger = logging.getLogger("koib.parsing")


@dataclass
class DocumentElement:
    content: str
    element_type: str
    source: str = ""
    page: int = 0
    heading: str = ""
    model: str = "unknown"
    element_id: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self) -> None:
        self.content = normalize_ocr_text(self.content)
        if not self.element_id:
            self.element_id = text_hash(f"{self.source}:{self.page}:{self.element_type}:{self.content[:300]}")

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    @property
    def is_structured(self) -> bool:
        return self.element_type in ("table", "formula", "figure")


def _expand_rect(rect: fitz.Rect, margin: float) -> fitz.Rect:
    return fitz.Rect(rect.x0 - margin, rect.y0 - margin, rect.x1 + margin, rect.y1 + margin)


def _is_scanned_page(page: fitz.Page, min_chars: int = OCR_MIN_TEXT_CHARS) -> bool:
    text = page.get_text("text").strip()
    if len(text) >= min_chars:
        return False
    images = page.get_images(full=True)
    if not images:
        return len(text) < min_chars

    page_area = max(1.0, page.rect.width * page.rect.height)
    for img_info in images:
        try:
            xref = img_info[0]
            base_image = page.parent.extract_image(xref)
            if not base_image:
                continue
            img = Image.open(io.BytesIO(base_image["image"]))
            try:
                if img.width * img.height / page_area > 0.8:
                    return True
            finally:
                img.close()
        except Exception:
            continue
    return False


def _ocr_image(image_pil: Image.Image, lang: str = "rus+eng") -> str:
    if image_pil is None:
        return ""
    try:
        import pytesseract

        text = normalize_ocr_text(pytesseract.image_to_string(image_pil, lang=lang, config="--psm 6"))
        if len(text) >= 30:
            return text
    except Exception as exc:
        logger.debug("Tesseract OCR error: %s", exc)
    return ""


def _extract_tables_from_page(page: fitz.Page) -> List[Dict[str, Any]]:
    tables: List[Dict[str, Any]] = []
    try:
        tab_finder = page.find_tables()
        for tab in tab_finder:
            try:
                rows = tab.extract()
                if not rows or len(rows) < 2:
                    continue
                num_cols = max(len(r) for r in rows) if rows else 0
                non_empty = sum(1 for row in rows for cell in row if cell and str(cell).strip())
                if non_empty < 2:
                    continue
                md_lines = []
                for i, row in enumerate(rows):
                    cells = [normalize_ocr_text(str(c)) if c else "" for c in row]
                    while len(cells) < num_cols:
                        cells.append("")
                    md_lines.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        md_lines.append("| " + " | ".join(["---"] * num_cols) + " |")
                markdown = "\n".join(md_lines)
                tables.append(
                    {
                        "text": markdown,
                        "num_rows": len(rows),
                        "num_cols": num_cols,
                        "bbox": tuple(tab.bbox) if hasattr(tab, "bbox") else (0, 0, 0, 0),
                    }
                )
            except Exception as exc:
                logger.debug("Ошибка конвертации таблицы: %s", exc)
    except AttributeError:
        logger.debug("PyMuPDF < 1.23: find_tables() недоступен")
    except Exception as exc:
        logger.debug("Ошибка поиска таблиц: %s", exc)
    return tables


def _detect_formulas_in_text(text: str) -> List[Dict[str, Any]]:
    """Strict formula extraction that does not treat '-' and '/' in normal words as formulas."""
    formulas: List[Dict[str, Any]] = []
    for match in _LATEX_BLOCK_RE.finditer(text):
        formulas.append(
            {"content": match.group(1).strip(), "formula_type": "latex_block", "start": match.start(), "end": match.end()}
        )
    for match in _LATEX_INLINE_RE.finditer(text):
        formulas.append(
            {"content": match.group(1).strip(), "formula_type": "latex_inline", "start": match.start(), "end": match.end()}
        )

    for line in text.split("\n"):
        line = normalize_ocr_text(line)
        if len(line) < 5 or len(line) > 260 or line.startswith("|"):
            continue
        if re.match(r"^\s*(рисунок|рис\.|таблица|схема)\s+", line, re.IGNORECASE):
            continue
        if _MATH_STRICT_RE.search(line) and _is_true_formula_text(line):
            start = text.find(line)
            if start < 0:
                start = 0
            already_found = any(f["start"] <= start <= f["end"] for f in formulas)
            if not already_found:
                formulas.append({"content": line, "formula_type": "suspected_formula", "start": start, "end": start + len(line)})
    return formulas


def _element_model(content: str, default_model: str) -> str:
    detected_model, conf = detect_model_in_text(content)
    return detected_model if conf > 0.3 else default_model


def parse_pdf(file_path: Path, model_hint: str = "") -> List[DocumentElement]:
    if not file_path.exists():
        logger.error("Файл не найден: %s", file_path)
        return []
    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    elements: List[DocumentElement] = []
    try:
        doc = fitz.open(str(file_path))
    except Exception as exc:
        logger.error("Не удалось открыть PDF %s: %s", filename, exc)
        return []

    logger.info("Парсинг PDF: %s (%s стр.)", filename, len(doc))

    # Документ-уровневая «липкая» детекция модели (Проблема 4): один проход по
    # тексту страниц закрепляет модель по титулам/децимальным кодам и переносит
    # её на отсканированные страницы без текста. Это устраняет model='unknown'
    # и смешивание интерфейсов КОИБ-2010/2017 в склеенных PDF.
    page_texts = [doc[i].get_text("text") for i in range(len(doc))]
    tagger = StickyModelTagger(default_model=model if model in {"koib2010", "koib2017a", "koib2017b"} else "unknown")
    page_models = tagger.assign_pages(page_texts)
    # запасной общий model для элементов, где страница так и осталась unknown
    detected_model, confidence = detect_model_in_text("\n".join(page_texts[:12]))
    fallback_model = detected_model if confidence > 0.3 else model

    current_heading = ""
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_model = page_models[page_num]
            if page_model not in {"koib2010", "koib2017a", "koib2017b"}:
                page_model = fallback_model
            page_text = page.get_text("text").strip()
            if _is_scanned_page(page):
                pix = page.get_pixmap(dpi=OCR_DPI)
                try:
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_text = deep_clean_ocr_text(_ocr_image(img))
                    if ocr_text:
                        elements.append(
                            DocumentElement(
                                content=ocr_text,
                                element_type="text",
                                source=filename,
                                page=page_num + 1,
                                model=page_model,
                                heading=current_heading,
                                metadata={"ocr": True},
                            )
                        )
                    img.close()
                finally:
                    del pix
                    gc.collect()
                continue

            headings = extract_headings(page_text)
            if headings:
                current_heading = headings[0]
                for heading in headings[:3]:
                    elements.append(
                        DocumentElement(
                            content=heading,
                            element_type="heading",
                            source=filename,
                            page=page_num + 1,
                            model=page_model,
                            heading=heading,
                        )
                    )

            for table_data in _extract_tables_from_page(page):
                table_text = normalize_ocr_text(table_data["text"])
                if table_text:
                    elements.append(
                        DocumentElement(
                            content=table_text,
                            element_type="table",
                            source=filename,
                            page=page_num + 1,
                            model=page_model,
                            heading=current_heading,
                            metadata={
                                "num_rows": table_data["num_rows"],
                                "num_cols": table_data["num_cols"],
                                "bbox": table_data["bbox"],
                            },
                        )
                    )

            for formula_data in _detect_formulas_in_text(page_text):
                elements.append(
                    DocumentElement(
                        content=formula_data["content"],
                        element_type="formula",
                        source=filename,
                        page=page_num + 1,
                        model=page_model,
                        heading=current_heading,
                        metadata={"formula_type": formula_data["formula_type"]},
                    )
                )

            for img_idx, img_info in enumerate(page.get_images(full=True)):
                img = None
                try:
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue
                    img = Image.open(io.BytesIO(base_image["image"]))
                    if img.width < MIN_IMAGE_WIDTH or img.height < MIN_IMAGE_HEIGHT:
                        continue
                    caption = find_figure_caption(page_text)
                    content = caption if caption else f"Изображение {img_idx + 1}"
                    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
                    img_filename = f"{file_path.stem}_p{page_num + 1}_img{img_idx}.png"
                    img_path = FIGURES_DIR / img_filename
                    img.save(str(img_path))
                    elements.append(
                        DocumentElement(
                            content=content,
                            element_type="figure",
                            source=filename,
                            page=page_num + 1,
                            model=page_model,
                            heading=current_heading,
                            metadata={
                                # Переносимый путь относительно OUTPUT_DIR
                                # («figures/..._img0.png») — см. config.rel_figure_path.
                                "image_path": rel_figure_path(img_path),
                                "width": img.width,
                                "height": img.height,
                                "figure_index": img_idx,
                            },
                        )
                    )
                except Exception as exc:
                    logger.debug("Ошибка изображения: %s", exc)
                finally:
                    if img is not None:
                        try:
                            img.close()
                        except Exception:
                            pass

            cleaned = deep_clean_ocr_text(page_text)
            if len(cleaned) >= OCR_MIN_TEXT_CHARS:
                elements.append(
                    DocumentElement(
                        content=cleaned,
                        element_type="text",
                        source=filename,
                        page=page_num + 1,
                        model=page_model,
                        heading=current_heading,
                    )
                )
            gc.collect()
    finally:
        doc.close()
        gc.collect()

    logger.info("Извлечено %s элементов из %s", len(elements), filename)
    return elements


def _iter_docx_blocks(doc) -> Iterable[Any]:
    if CT_P is None or CT_Tbl is None:
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            yield table
        return
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, doc)


def _docx_table_to_markdown(table) -> tuple[str, int, int]:
    rows_data = [[normalize_ocr_text(cell.text) for cell in row.cells] for row in table.rows]
    rows_data = [row for row in rows_data if any(cell.strip() for cell in row)]
    if not rows_data:
        return "", 0, 0
    num_cols = max(len(r) for r in rows_data)
    header = list(rows_data[0])
    while len(header) < num_cols:
        header.append(f"Кол.{len(header) + 1}")
    md_lines = ["| " + " | ".join(header[:num_cols]) + " |", "| " + " | ".join(["---"] * num_cols) + " |"]
    for row in rows_data[1:]:
        row = list(row)
        while len(row) < num_cols:
            row.append("")
        md_lines.append("| " + " | ".join(row[:num_cols]) + " |")
    return "\n".join(md_lines), len(rows_data), num_cols


def parse_docx(file_path: Path, model_hint: str = "") -> List[DocumentElement]:
    if not file_path.exists():
        logger.error("Файл не найден: %s", file_path)
        return []
    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    elements: List[DocumentElement] = []
    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        logger.error("Не удалось открыть DOCX %s: %s", filename, exc)
        return []

    logger.info("Парсинг DOCX: %s", filename)
    current_heading = ""
    text_buffer: List[str] = []
    text_block_index = 0

    def flush_text() -> None:
        nonlocal text_buffer, text_block_index
        if not text_buffer:
            return
        combined = normalize_ocr_text("\n".join(text_buffer))
        text_buffer = []
        if len(combined) < 50:
            return
        text_block_index += 1
        elements.append(
            DocumentElement(
                content=combined,
                element_type="text",
                source=filename,
                page=text_block_index,
                model=_element_model(combined, model),
                heading=current_heading,
                metadata={"docx_block": text_block_index, "page_is_block": True},
            )
        )

    table_index = 0
    for block in _iter_docx_blocks(doc):
        if Paragraph is not None and isinstance(block, Paragraph):
            text = normalize_ocr_text(block.text)
            if not text:
                continue
            style_name = block.style.name.lower() if block.style else ""
            is_heading = "heading" in style_name or "заголовок" in style_name or bool(extract_headings(text[:200]))
            if is_heading and len(text) <= 180:
                flush_text()
                current_heading = text
                elements.append(
                    DocumentElement(
                        content=text,
                        element_type="heading",
                        source=filename,
                        page=text_block_index + 1,
                        model=_element_model(text, model),
                        heading=text,
                        metadata={"docx_heading": True},
                    )
                )
                continue
            text_buffer.append(text)
            if sum(len(p) for p in text_buffer) >= 8000:
                flush_text()
        elif DocxTable is not None and isinstance(block, DocxTable):
            flush_text()
            table_md, rows, cols = _docx_table_to_markdown(block)
            if table_md:
                table_index += 1
                elements.append(
                    DocumentElement(
                        content=table_md,
                        element_type="table",
                        source=filename,
                        page=max(text_block_index, 1),
                        model=_element_model(table_md, model),
                        heading=current_heading,
                        metadata={"num_rows": rows, "num_cols": cols, "table_index": table_index},
                    )
                )

    flush_text()
    gc.collect()
    logger.info("Извлечено %s элементов из %s", len(elements), filename)
    return elements


def _read_text_file(file_path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return file_path.read_text(encoding=enc, errors="strict")
        except UnicodeDecodeError:
            continue
    return file_path.read_text(encoding="utf-8", errors="replace")


def _text_blocks(text: str, max_chars: int = 12000) -> Iterable[str]:
    text = normalize_ocr_text(text)
    parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    buf: List[str] = []
    size = 0
    for part in parts or [text]:
        if size + len(part) > max_chars and buf:
            yield "\n\n".join(buf)
            buf, size = [], 0
        buf.append(part)
        size += len(part)
    if buf:
        yield "\n\n".join(buf)


def parse_text_file(file_path: Path, model_hint: str = "") -> List[DocumentElement]:
    if _is_supported_artifact(file_path) or looks_like_jsonl_artifact(file_path):
        logger.warning("Файл %s похож на артефакт индекса. Используйте --ingest-artifacts/--artifacts-dir, не data/docs.", file_path.name)
        return []
    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    raw = _read_text_file(file_path)
    elements: List[DocumentElement] = []
    for i, block in enumerate(_text_blocks(raw), 1):
        if len(block) < 50:
            continue
        elements.append(
            DocumentElement(
                content=block,
                element_type="text",
                source=filename,
                page=i,
                model=_element_model(block, model),
                metadata={"text_block": i, "recognized_text": True},
            )
        )
    return elements


def _sniff_delimiter(sample: str) -> str:
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except Exception:
        return ","


def parse_csv_file(file_path: Path, model_hint: str = "") -> List[DocumentElement]:
    if looks_like_jsonl_artifact(file_path):
        logger.warning("CSV %s похож на JSONL-артефакт индекса. Используйте импорт артефактов.", file_path.name)
        return []
    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    raw = _read_text_file(file_path)
    sample = raw[:8192]
    delimiter = _sniff_delimiter(sample)
    rows: List[List[str]] = []
    try:
        reader = csv.reader(raw.splitlines(), delimiter=delimiter)
        rows = [[normalize_ocr_text(cell) for cell in row] for row in reader]
    except Exception:
        rows = []

    col_counts = [len(r) for r in rows[:300] if r]
    multi_col_ratio = sum(1 for c in col_counts if c > 1) / max(1, len(col_counts))
    if rows and multi_col_ratio >= 0.35:
        elements: List[DocumentElement] = []
        batch_size = 80
        for start in range(0, len(rows), batch_size):
            batch = [r for r in rows[start : start + batch_size] if any(cell.strip() for cell in r)]
            if not batch:
                continue
            num_cols = max(len(r) for r in batch)
            header = list(batch[0])
            while len(header) < num_cols:
                header.append(f"Кол.{len(header) + 1}")
            md_lines = ["| " + " | ".join(header[:num_cols]) + " |", "| " + " | ".join(["---"] * num_cols) + " |"]
            for row in batch[1:]:
                row = list(row)
                while len(row) < num_cols:
                    row.append("")
                md_lines.append("| " + " | ".join(row[:num_cols]) + " |")
            table_md = "\n".join(md_lines)
            elements.append(
                DocumentElement(
                    content=table_md,
                    element_type="table",
                    source=filename,
                    page=start // batch_size + 1,
                    model=_element_model(table_md, model),
                    metadata={"num_rows": len(batch), "num_cols": num_cols, "csv_table": True},
                )
            )
        return elements

    # Most OCR "CSV" exports are actually a single-column recognized text stream.
    return parse_text_file(file_path, model_hint=model)


def parse_document(file_path: Path, engine: Optional[str] = None, model_hint: str = "") -> List[DocumentElement]:
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path, model_hint)
    if ext == ".docx":
        return parse_docx(file_path, model_hint)
    if ext == ".doc":
        logger.error("Legacy .doc не поддерживается python-docx. Конвертируйте файл в .docx или PDF: %s", file_path)
        return []
    if ext == ".csv":
        return parse_csv_file(file_path, model_hint)
    if ext in {".txt", ".md"}:
        return parse_text_file(file_path, model_hint)
    return []
